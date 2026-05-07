"""Generates DOCX report from ResearchReport model."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from src.logging_config import get_logger
from src.models.common import MaturityStage
from src.models.report import ResearchReport

logger = get_logger("reporters.docx")


class DocxReporter:
    """Renders a ResearchReport into a DOCX file."""

    def generate(self, report: ResearchReport, output_path: Path | str) -> Path:
        """Generate DOCX report and write it to disk."""
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                'DOCX output requires python-docx. Install it with: pip install -e ".[docx]"'
            ) from exc

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        document = Document()
        self._configure_styles(document)
        self._render(report, document)
        document.save(output_path)

        logger.info("DOCX report written to %s", output_path)
        return output_path

    def _configure_styles(self, document: Any) -> None:
        """Apply light, readable defaults to the generated document."""
        from docx.shared import Pt

        normal = document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)

        for style_name, size in [("Title", 22), ("Heading 1", 16), ("Heading 2", 13)]:
            style = document.styles[style_name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)

    def _render(self, report: ResearchReport, document: Any) -> None:
        title = report.title or "Product Research Report"
        document.add_heading(title, level=0)
        self._add_metadata(report, document)
        self._add_section(document, "Executive Summary", report.executive_summary)
        self._add_research_method(report, document)
        self._add_decision_matrix(report, document)
        self._add_key_claims(report, document)
        self._add_technology_landscape(report, document)
        self._add_technology_relationships(report, document)
        self._add_path_deep_analysis(report, document)
        self._add_cross_analysis(report, document)
        self._add_maturity_map(report, document)
        self._add_workflows(report, document)
        self._add_industry_findings(report, document)
        self._add_academic_findings(report, document)
        self._add_engineering_analysis(report, document)
        self._add_reputation(report, document)
        self._add_feasibility(report, document)
        self._add_confidence_and_gaps(report, document)
        recommendations = report.recommended_strategy or report.recommendations
        if report.recommended_strategy and report.recommendations:
            recommendations = f"{report.recommended_strategy}\n\n{report.recommendations}"
        self._add_section(document, "Recommendations", recommendations)
        self._add_sources(report, document)

    def _add_metadata(self, report: ResearchReport, document: Any) -> None:
        rows = [
            ("Generated", report.generated_at),
            ("Session", report.session_id),
            ("Input", report.original_input),
        ]
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for key, value in rows:
            if not value:
                continue
            cells = table.add_row().cells
            cells[0].text = key
            cells[1].text = value

    def _add_section(self, document: Any, heading: str, text: str) -> None:
        if not text:
            return
        document.add_heading(heading, level=1)
        document.add_paragraph(text)

    def _add_research_method(self, report: ResearchReport, document: Any) -> None:
        if not report.research_questions and not report.methodology_summary:
            return
        document.add_heading("Research Question and Method", level=1)
        for question in report.research_questions:
            document.add_paragraph(question, style="List Bullet")
        if report.methodology_summary:
            document.add_paragraph(report.methodology_summary)

    def _add_decision_matrix(self, report: ResearchReport, document: Any) -> None:
        if not report.decision_matrix:
            return
        document.add_heading("Decision Matrix", level=1)
        headers = [
            "Option",
            "User value",
            "Feasibility",
            "Maturity",
            "Ecosystem",
            "Cost/Risk",
            "Evidence",
            "Verdict",
        ]
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        for row in report.decision_matrix:
            values = [
                row.option or row.path_id,
                row.user_value,
                row.technical_feasibility,
                row.maturity,
                row.ecosystem_strength,
                row.cost_risk,
                row.evidence_strength,
                row.verdict,
            ]
            cells = table.add_row().cells
            for i, value in enumerate(values):
                cells[i].text = value

    def _add_key_claims(self, report: ResearchReport, document: Any) -> None:
        if not report.key_claims:
            return
        document.add_heading("Key Claims and Evidence", level=1)
        for claim in report.key_claims:
            document.add_heading(claim.claim, level=2)
            if claim.confidence:
                document.add_paragraph(f"Confidence: {claim.confidence}")
            self._add_inline_list(
                document,
                "Supporting evidence",
                claim.supporting_evidence,
            )
            self._add_inline_list(
                document,
                "Contradicting evidence or caveats",
                claim.contradicting_evidence,
            )
            if claim.implication:
                document.add_paragraph(f"Implication: {claim.implication}")
            self._add_inline_list(document, "Sources", claim.source_urls)

    def _add_technology_landscape(self, report: ResearchReport, document: Any) -> None:
        if not report.technology_landscape:
            return
        document.add_heading("Technology Landscape", level=1)
        headers = [
            "Technology",
            "Category",
            "Maturity",
            "Industry",
            "Academic",
            "Engineering",
            "Reputation",
        ]
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        for tech in report.technology_landscape:
            cells = table.add_row().cells
            values = [
                tech.name,
                tech.category,
                tech.maturity.value,
                f"{tech.industry_score:.1f}",
                f"{tech.academic_score:.1f}",
                f"{tech.engineering_score:.1f}",
                f"{tech.reputation_score:.1f}",
            ]
            for i, value in enumerate(values):
                cells[i].text = value
        for tech in report.technology_landscape:
            if tech.description:
                paragraph = document.add_paragraph()
                paragraph.add_run(f"{tech.name}: ").bold = True
                paragraph.add_run(tech.description)

    def _add_maturity_map(self, report: ResearchReport, document: Any) -> None:
        if not report.maturity_assessments:
            return
        document.add_heading("Technology Maturity Map", level=1)
        stage_labels = {
            MaturityStage.EARLY_PROTOTYPE: "Early Prototype",
            MaturityStage.DEVELOPMENT: "In Development",
            MaturityStage.MATURE: "Mature",
            MaturityStage.CUTTING_EDGE: "Cutting Edge",
            MaturityStage.ACADEMIC_FRONTIER: "Academic Frontier",
        }
        for assessment in report.maturity_assessments:
            document.add_heading(f"Path: {assessment.path_id}", level=2)
            if assessment.overall_maturity_summary:
                document.add_paragraph(assessment.overall_maturity_summary)
            for tech in assessment.technologies:
                label = stage_labels.get(tech.stage, tech.stage.value)
                trend = f", {tech.trend}" if tech.trend else ""
                document.add_paragraph(
                    f"{tech.technology} [{label}{trend}]: {tech.evidence}",
                    style="List Bullet",
                )

    def _add_workflows(self, report: ResearchReport, document: Any) -> None:
        if not report.workflows:
            return
        document.add_heading("Implementation Workflows", level=1)
        for index, workflow in enumerate(report.workflows, 1):
            document.add_heading(f"Workflow {index}: {workflow.name}", level=2)
            if workflow.description:
                document.add_paragraph(workflow.description)
            for step in workflow.steps:
                technologies = ", ".join(step.technologies)
                suffix = f" Technologies: {technologies}." if technologies else ""
                considerations = (
                    f" Considerations: {step.considerations}" if step.considerations else ""
                )
                document.add_paragraph(
                    f"{step.step}. {step.description}.{suffix}{considerations}",
                    style="List Number",
                )
            self._add_inline_list(document, "Pros", workflow.pros)
            self._add_inline_list(document, "Cons", workflow.cons)

    def _add_path_deep_analysis(self, report: ResearchReport, document: Any) -> None:
        if not report.path_deep_analysis:
            return
        document.add_heading("研究路线深度分析", level=1)
        for pa in report.path_deep_analysis:
            document.add_heading(pa.title, level=2)
            if pa.technical_overview:
                document.add_paragraph(pa.technical_overview)
            if pa.key_technologies_detail:
                document.add_heading("核心技术详解", level=3)
                for td in pa.key_technologies_detail:
                    p = document.add_paragraph()
                    p.add_run(f"{td.name}\n").bold = True
                    for label, value in [
                        ("是什么", td.what_it_is), ("工作原理", td.how_it_works),
                        ("优势", "；".join(td.pros)), ("劣势", "；".join(td.cons)),
                        ("实现建议", td.implementation_notes),
                        ("产业证据", td.industry_evidence),
                        ("学术证据", td.academic_evidence),
                        ("工程证据", td.engineering_evidence),
                    ]:
                        if value:
                            document.add_paragraph(f"{label}：{value}", style="List Bullet")
            if pa.cross_references:
                document.add_heading("交叉关联", level=3)
                document.add_paragraph(pa.cross_references)
            pcs = pa.pros_cons_summary
            if pcs.strengths or pcs.weaknesses:
                document.add_heading("优劣势总结", level=3)
                self._add_inline_list(document, "核心优势", pcs.strengths)
                self._add_inline_list(document, "核心劣势", pcs.weaknesses)
                if pcs.best_for:
                    document.add_paragraph(f"最适合场景：{pcs.best_for}")
                if pcs.not_suitable_for:
                    document.add_paragraph(f"不适合场景：{pcs.not_suitable_for}")

    def _add_cross_analysis(self, report: ResearchReport, document: Any) -> None:
        ca = report.cross_analysis
        if not ca.industry_academic_alignment and not ca.academic_engineering_gap:
            return
        document.add_heading("交叉分析：产业-学术-工程的互证与矛盾", level=1)
        if ca.industry_academic_alignment:
            document.add_heading("产业与学术的一致性", level=2)
            document.add_paragraph(ca.industry_academic_alignment)
        if ca.academic_engineering_gap:
            document.add_heading("学术研究与工程实现的差距", level=2)
            document.add_paragraph(ca.academic_engineering_gap)
        if ca.evidence_quality_overview:
            document.add_heading("证据质量总览", level=2)
            document.add_paragraph(ca.evidence_quality_overview)
        if ca.key_contradictions:
            document.add_heading("关键矛盾与分歧", level=2)
            for kc in ca.key_contradictions:
                document.add_paragraph(kc, style="List Bullet")

    def _add_technology_relationships(self, report: ResearchReport, document: Any) -> None:
        tr = report.technology_relationships
        has_rel = (tr.complementary_pairs or tr.alternatives or tr.dependency_chains)
        if not has_rel:
            return
        document.add_heading("技术关系图谱", level=1)
        if tr.complementary_pairs:
            document.add_heading("互补技术", level=2)
            for pair in tr.complementary_pairs:
                text = f"{pair[0]} + {pair[1]}" if len(pair) >= 2 else str(pair)
                if len(pair) >= 3:
                    text += f"：{pair[2]}"
                document.add_paragraph(text, style="List Bullet")
        if tr.alternatives:
            document.add_heading("替代关系", level=2)
            for pair in tr.alternatives:
                text = f"{pair[0]} ↔ {pair[1]}" if len(pair) >= 2 else str(pair)
                if len(pair) >= 3:
                    text += f"：{pair[2]}"
                document.add_paragraph(text, style="List Bullet")
        if tr.dependency_chains:
            document.add_heading("依赖链", level=2)
            for chain in tr.dependency_chains:
                text = f"{chain[0]} → {chain[1]}" if len(chain) >= 2 else str(chain)
                if len(chain) >= 3:
                    text += f"：{chain[2]}"
                document.add_paragraph(text, style="List Bullet")

    def _add_industry_findings(self, report: ResearchReport, document: Any) -> None:
        if not report.industry_findings:
            return
        document.add_heading("Industry Findings", level=1)
        for finding in report.industry_findings:
            document.add_heading(f"Research Path: {finding.path_id}", level=2)
            for product in finding.products:
                document.add_paragraph(
                    f"{product.name} by {product.company}: {product.description}",
                    style="List Bullet",
                )
            for repo in finding.repos:
                document.add_paragraph(
                    f"{repo.name} - {repo.stars} stars, {repo.language}. {repo.repo_url}",
                    style="List Bullet",
                )
            for blog in finding.blog_summaries:
                document.add_paragraph(f"{blog.title}: {blog.url}", style="List Bullet")
                for point in blog.key_points[:3]:
                    document.add_paragraph(point, style="List Bullet 2")

    def _add_academic_findings(self, report: ResearchReport, document: Any) -> None:
        if not report.academic_findings:
            return
        document.add_heading("Academic Research", level=1)
        for finding in report.academic_findings:
            document.add_heading(f"Research Path: {finding.path_id}", level=2)
            for paper in finding.papers:
                venue = f" @ {paper.venue}" if paper.venue else ""
                cite_info = f", {paper.citation_count} citations" if paper.citation_count else ""
                document.add_paragraph(
                    f"{paper.title} ({paper.year}{venue}{cite_info})",
                    style="List Bullet",
                )
                self._add_inline_list(document, "Authors", paper.authors[:5])
                for label, value in [
                    ("Principles", paper.principles),
                    ("Methods", paper.methods),
                    ("Conclusions", paper.conclusions),
                    ("Limitations", paper.deficiencies),
                ]:
                    if value:
                        document.add_paragraph(f"{label}: {value}")

    def _add_engineering_analysis(self, report: ResearchReport, document: Any) -> None:
        if not report.engineering_analyses:
            return
        document.add_heading("Engineering Analysis", level=1)
        for analysis in report.engineering_analyses:
            document.add_heading(f"Path: {analysis.path_id}", level=2)
            deployment = analysis.deployment_assessment
            if deployment.deployment_complexity:
                document.add_paragraph(f"Deployment complexity: {deployment.deployment_complexity}")
            self._add_inline_list(
                document,
                "Infrastructure",
                deployment.infrastructure_requirements,
            )
            self._add_inline_list(document, "Risks", deployment.risks)
            if analysis.implementation_recommendations:
                document.add_paragraph(
                    f"Recommendations: {analysis.implementation_recommendations}"
                )
            self._add_inline_list(
                document,
                "Recommended stack",
                analysis.technology_stack_recommendation,
            )

    def _add_reputation(self, report: ResearchReport, document: Any) -> None:
        reputation = report.reputation_report
        if not reputation.companies and not reputation.labs and not reputation.technology_scores:
            return
        document.add_heading("Reputation and Credibility", level=1)
        if reputation.summary:
            document.add_paragraph(reputation.summary)
        for company in reputation.companies:
            document.add_paragraph(
                f"{company.company_name} ({company.category}): "
                f"score={company.reputation_score:.2f}, sentiment={company.user_sentiment}",
                style="List Bullet",
            )
        for lab in reputation.labs:
            document.add_paragraph(
                f"{lab.lab_name} at {lab.institution}: score={lab.prestige_score:.2f}",
                style="List Bullet",
            )
        for score in reputation.technology_scores:
            document.add_paragraph(
                f"{score.name}: overall credibility {score.overall_score:.2f}",
                style="List Bullet",
            )

    def _add_feasibility(self, report: ResearchReport, document: Any) -> None:
        if not report.feasibility_assessments:
            return
        document.add_heading("Feasibility Assessment", level=1)
        for assessment in report.feasibility_assessments:
            verdict = "Recommended" if assessment.recommended else "Not Recommended"
            document.add_heading(f"{assessment.path_title} ({verdict})", level=2)
            document.add_paragraph(
                f"Feasibility score: {assessment.overall_feasibility:.0%}"
            )
            document.add_paragraph(f"Technical risk: {assessment.technical_risk}")
            if assessment.resource_requirements:
                document.add_paragraph(f"Resources: {assessment.resource_requirements}")
            if assessment.time_estimate:
                document.add_paragraph(f"Time estimate: {assessment.time_estimate}")
            if assessment.rationale:
                document.add_paragraph(f"Rationale: {assessment.rationale}")
            self._add_inline_list(document, "Key challenges", assessment.key_challenges)

    def _add_confidence_and_gaps(self, report: ResearchReport, document: Any) -> None:
        if not report.confidence_assessment and not report.evidence_gaps and not report.assumptions:
            return
        document.add_heading("Confidence, Gaps, and Assumptions", level=1)
        if report.confidence_assessment:
            document.add_paragraph(report.confidence_assessment)
        self._add_inline_list(document, "Evidence gaps", report.evidence_gaps)
        self._add_inline_list(document, "Assumptions", report.assumptions)

    def _add_sources(self, report: ResearchReport, document: Any) -> None:
        if not report.all_sources:
            return
        document.add_heading("Sources", level=1)
        for source in report.all_sources:
            source_type = source.source_type.value if source.source_type else "web"
            document.add_paragraph(
                f"{source.title} [{source_type}] - {source.url}",
                style="List Bullet",
            )

    def _add_inline_list(self, document: Any, label: str, values: list[str]) -> None:
        if values:
            document.add_paragraph(f"{label}: {', '.join(values)}")
