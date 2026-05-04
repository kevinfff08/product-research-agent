"""Tests for DocxReporter."""

from __future__ import annotations

import pytest

docx = pytest.importorskip("docx")

from docx import Document

from src.models.common import MaturityStage, SourceReference, SourceType
from src.models.report import ResearchReport, TechnologyEntry, DecisionMatrixRow, EvidenceClaim
from src.reporters.docx_reporter import DocxReporter


@pytest.fixture
def docx_report() -> ResearchReport:
    return ResearchReport(
        title="Tech Landscape: AI Code Review",
        session_id="test-123",
        generated_at="2026-05-04T00:00:00Z",
        original_input="AI code review tool",
        executive_summary="This report analyzes the AI code review landscape.",
        research_questions=["Should we build or reuse an AI code review stack?"],
        methodology_summary="Compared products, papers, repositories, and maturity signals.",
        decision_matrix=[
            DecisionMatrixRow(
                option="Static Analysis + LLM",
                path_id="p1",
                user_value="high",
                technical_feasibility="high",
                maturity="mature",
                ecosystem_strength="strong",
                cost_risk="medium",
                evidence_strength="moderate",
                verdict="recommended",
            ),
        ],
        key_claims=[
            EvidenceClaim(
                claim="Parser-grounded review is practical.",
                supporting_evidence=["tree-sitter is mature."],
                confidence="medium",
            ),
        ],
        technology_landscape=[
            TechnologyEntry(
                name="tree-sitter",
                category="backend",
                maturity=MaturityStage.MATURE,
                description="Fast parser",
                industry_score=0.9,
                academic_score=0.5,
                engineering_score=0.95,
                reputation_score=0.85,
            ),
        ],
        recommendations="Start with the simple approach.",
        all_sources=[
            SourceReference(
                url="https://example.com",
                title="Source 1",
                source_type=SourceType.WEB,
            ),
        ],
    )


def test_generate_creates_docx(tmp_path, docx_report):
    output = tmp_path / "report.docx"
    result = DocxReporter().generate(docx_report, output)

    assert result.exists()
    document = Document(result)
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Tech Landscape: AI Code Review" in text
    assert "Executive Summary" in text
    assert "Research Question and Method" in text
    assert "Key Claims and Evidence" in text
    assert "Start with the simple approach." in text
    assert document.tables


def test_empty_report(tmp_path):
    output = tmp_path / "empty.docx"
    result = DocxReporter().generate(ResearchReport(title="Empty"), output)

    document = Document(result)
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Empty" in text
